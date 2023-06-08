import os
import pandas as pd
from docx import Document
from docx2pdf import convert
import tkinter as tk
from tkinter import messagebox
import tempfile
from concurrent.futures import ThreadPoolExecutor
import pythoncom


def replace_fields_in_document(document, first_name, last_name, username, phone_number, mod, selected_users):
    #this section is for replacing in the first paragraph
    for paragraph in document.paragraphs:
        if "{F_Name}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{F_Name}", first_name)

    #this section is for replacing in the table
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{F_Name}" in cell.text:
                    cell.text = cell.text.replace("{F_Name}", first_name)
                if "{L_Name}" in cell.text:
                    cell.text = cell.text.replace("{L_Name}", last_name)
                if "{Username}" in cell.text:
                    cell.text = cell.text.replace("{Username}", username)
                if "{Phone_Number}" in cell.text:
                    cell.text = cell.text.replace("{Phone_Number}", phone_number)
                if "{MOD}" in cell.text:
                    cell.text = cell.text.replace("{MOD}", str(mod))


#loads template document, calls replace fields and then saves the modified doc as a temp doc
#creates a pdf with the doc and uses the COM library to be able to do this
#removes the temporary word doc because it has been converted to a pdf
def create_pdf(starting_date, first_name, last_name, username, phone_number, mod, template_file_path, output_pdf_path, selected_users):
    # Load the template document
    template_doc = Document(template_file_path)
    print("created")

    # Replace the fields in the document with the name, username, phone number, and mod
    replace_fields_in_document(template_doc, first_name, last_name, username, phone_number, mod, selected_users)
    print("fields replaced")

    # Save the modified document as a temporary file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        temp_file_path = temp_file.name
        template_doc.save(temp_file_path)

    # Generate the output PDF path with the username
    output_pdf_dir = os.path.dirname(output_pdf_path)
    output_pdf_filename = os.path.basename(output_pdf_path)
    output_pdf_filename = output_pdf_filename.replace("{Username}", username)
    output_pdf_path = os.path.join(output_pdf_dir, output_pdf_filename)
    print("pdf created")

    # Initialize the COM library
    pythoncom.CoInitialize()

    try:
        # Convert the Word document to PDF
        convert(temp_file_path, output_pdf_path)
    finally:
        # Uninitialize the COM library
        pythoncom.CoUninitialize()

    # Remove the temporary file
    os.remove(temp_file_path)
    print("file removed")


def create_pdfs(excel_file_path, template_file_path, selected_users):
    try:
        # Read Excel sheet and retrieve names
        data_frame = pd.read_excel(excel_file_path, engine='openpyxl')
        excel_values = data_frame[['Starting Date', 'First Name', 'Last Name', 'MOD', 'Phone Number']].fillna('').values


        # Create the Welcome PDFs for selected users
        for value in excel_values:
            starting_date, first_name, last_name, mod, phone_number = value

            if mod:
                mod = int(mod)
            else:
                mod = ''

            # Generate the username
            username = f"{first_name}.{last_name}{mod}"

            # Generate the output PDF path
            output_pdf_path = f"NPS_NH_Form-{starting_date}-{username}.pdf"

            # Check if the user is selected
            print(selected_users)

            # Create a thread pool
            executor = ThreadPoolExecutor()

            if selected_users.get(f"{first_name}.{last_name}{mod}") is not None:
                print("User Found in selected_users:", f"{first_name}.{last_name}{mod}")
                # Submit PDF creation task to the thread pool
                executor.submit(create_pdf, starting_date, first_name, last_name, username, phone_number,
                                mod, template_file_path, output_pdf_path, selected_users)
            else:
                print("User Not Found in selected_users:", f"{first_name}.{last_name}{mod}")

        executor.shutdown()

        messagebox.showinfo("PDF Creation Complete", "PDFs have been created successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")


# Create the GUI window
window = tk.Tk()
window.title("Welcome Sheet Creator")

# Set window dimensions and position
window_width = 400
window_height = 400
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
x = (screen_width / 2) - (window_width / 2)
y = (screen_height / 2) - (window_height / 2)
window.geometry(f"{int(window_width)}x{int(window_height)}+{int(x)}+{int(y)}")

# Create a frame to hold the widgets
frame = tk.Frame(window, padx=20, pady=20)
frame.pack()

# Create a label
label = tk.Label(frame, text="Welcome PDF Creator", font=("Arial", 16))
label.pack(pady=10)

excel_file_path = "New_Hires.xlsx"
template_file_path = "welcomeTemplateV2.docx"
selected_users = {}  # To store the selected user names


# Function to trigger PDF creation
def trigger_pdf_creation():
    global excel_file_path, template_file_path, selected_users

    if not template_file_path:
        messagebox.showerror("Error", "Please select a template document.")
        return

    try:
        create_pdfs(excel_file_path, template_file_path, selected_users)
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")


# Create a button to trigger PDF creation
create_pdf_button = tk.Button(frame, text="Create PDFs", command=trigger_pdf_creation, width=20)
create_pdf_button.pack(pady=20)


# Function to toggle user selection
def toggle_user_selection(user_name):
    if selected_users.get(user_name):
        selected_users.pop(user_name)
    else:
        selected_users[user_name] = True


# Function to update selected users
def update_selected_users():
    selected_user_names = [user_name for user_name, selected in selected_users.items() if selected]
    selected_users_label.config(text="Selected Users: " + ', '.join(selected_user_names))
    print(selected_user_names)


# Create a label for selected users
selected_users_label = tk.Label(frame, text="Selected Users: ")
selected_users_label.pack()


# Function to handle checkbox selection
def checkbox_selected(user_name):
    toggle_user_selection(user_name)
    update_selected_users()


# Create checkboxes for each user
def create_checkboxes():
    global selected_users

    selected_users = {}
    checkboxes_frame = tk.Frame(frame)
    checkboxes_frame.pack()

    try:
        data_frame = pd.read_excel(excel_file_path, engine='openpyxl')
        users = data_frame[['First Name', 'Last Name', 'MOD', 'Phone Number']].fillna('').values.tolist()

        for user in users:
            first_name, last_name, mod, phone_number = user

            if mod:
                mod = int(mod)
            else:
                mod = ''

            user_name = f"{first_name}.{last_name}{mod}"

            var = tk.BooleanVar()

            checkbox = tk.Checkbutton(checkboxes_frame, text=user_name, variable=var,
                                      command=lambda name=user_name: checkbox_selected(name))
            checkbox.pack(anchor=tk.W)

        update_selected_users()
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred while reading the Excel file: {str(e)}")


# Call the function to create checkboxes initially
create_checkboxes()

window.mainloop()
